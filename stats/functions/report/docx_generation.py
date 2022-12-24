import pandas as pd
from docx import Document

def docx_generation(new_species_maillettes, new_species_bossy, new_species_isole, rare_species_maillettes: pd.DataFrame, rare_species_bossy: pd.DataFrame, rare_species_isole: pd.DataFrame, total_species_maillettes, total_species_bossy, total_species_isole, year):

    document = Document()

    pre_title = document.add_paragraph('')
    pre_title.add_run("Observation et évolution des espèces fongiques sur les troncs du chemin des Maillettes et de ceux de la route du Pont de Bossy").bold = True

    document.add_paragraph('')

    document.add_heading("Rapport intermédiaire " + year, level=0)

    data_title = document.add_paragraph('')
    data_title.add_run('Quelques données pour le rapport :').bold = True

    document.add_paragraph("Nombre total d'espèces recensées au chemin des Maillettes: " + str(total_species_maillettes))
    document.add_paragraph("Nombre total d'espèces recensées au chemin du Pont-de-Bossy: " + str(total_species_bossy))
    document.add_paragraph("Nombre total d'espèces recensées du tronc isolé au chemin du Pont-de-Bossy: " + str(total_species_isole))

    document.add_paragraph('')

    introduction_title = document.add_paragraph('')
    introduction_title.add_run('Introduction :').bold = True

    document.add_paragraph('De nombreux champignons de la liste rouge des espèces menacées en Suisse croissent sur de vieux arbres morts, de grands diamètres, principalement de feuillus. De tels arbres sont rarement rencontrés en forêt et par conséquent une niche écologique manque souvent pour ces espèces. Afin de créer un biotope favorable pour ces champignons, des troncs de feuillus de grands diamètres (chênes et peupliers principalement, mais aussi de saule et de marronniers) ont été déposés au sol, en bordure d’une route, en lisière de forêt, sur un sol humide.')
    document.add_paragraph('Un suivi mycologique régulier a été entrepris dès 2014.')

    document.add_paragraph('')

    material_title = document.add_paragraph('')
    material_title.add_run('Matériel et méthode :').bold = True

    document.add_paragraph('')

    result_title = document.add_paragraph('')
    result_title.add_run('Résultat ').bold = True
    result_title.add_run('(au 31.12.' + year + ')')

    document.add_paragraph('')

    material_title = document.add_paragraph('')
    material_title.add_run('Discussion :').bold = True

    document.add_paragraph('')

    annexes = document.add_paragraph('')
    annexes.add_run('Annexes :').bold = True

    document.add_paragraph('Tableau 1 : ' + str(new_species_maillettes.shape[0]) + ' nouvelles espèces trouvées en ' + year + ' sur les troncs du chemin des Maillettes')
    document.add_paragraph('Tableau 2 : ' + str(new_species_bossy.shape[0]) + ' nouvelles espèces trouvées en ' + year + ' sur les troncs du chemin du Pont-de-Bossy')
    document.add_paragraph('Tableau 3 : ' + str(new_species_isole.shape[0]) + ' nouvelles espèces trouvées en ' + year + ' sur le tronc SMG du chemin du Pont-de-Bossy')
    document.add_paragraph('Tableau 4 : ' + str(rare_species_maillettes.shape[0]) + ' espèces de la liste rouge, rares ou assez rares trouvées en ' + year + ' sur les troncs du chemin des Maillettes')
    document.add_paragraph('Tableau 5 : ' + str(rare_species_bossy.shape[0]) + ' espèces de la liste rouge, rares ou assez rares trouvées en ' + year + ' sur les troncs du chemin du Pont-de-Bossy')
    document.add_paragraph('Tableau 6 : ' + str(rare_species_isole.shape[0]) + ' espèces de la liste rouge, rares ou assez rares trouvées en ' + year + ' sur le tronc SMG du chemin du Pont-de-Bossy')

    document.add_page_break()

    annexe1 = document.add_paragraph('')
    annexe1.add_run(
        'Tableau 1 : ' + str(new_species_maillettes.shape[0]) + ' nouvelles espèces trouvées en ' + year + ' sur les troncs du chemin des Maillettes').bold = True

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Espèce'
    hdr_cells[1].text = 'Statut liste rouge'
    hdr_cells[2].text = 'Troncs'
    for index, row in new_species_maillettes.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Espèce"])
        if row["Espèce actuelle"] != row["Espèce"]:
            row_cells[0].text = str(row["Espèce"] + " (= " + row["Espèce actuelle"] + ")")
        row_cells[1].text = str(row["Liste rouge"])
        row_cells[2].text = str(row["Tronc"])

    document.add_page_break()

    annexe2 = document.add_paragraph('')
    annexe2.add_run(
        'Tableau 2 : ' + str(new_species_bossy.shape[0]) + ' nouvelles espèces trouvées en ' + year + ' sur les troncs du chemin du Pont-de-Bossy').bold = True

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Espèce'
    hdr_cells[1].text = 'Statut liste rouge'
    hdr_cells[2].text = 'Troncs'
    for index, row in new_species_bossy.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Espèce"])
        if row["Espèce actuelle"] != row["Espèce"]:
            row_cells[0].text = str(row["Espèce"] + " (= " + row["Espèce actuelle"] + ")")
        row_cells[1].text = str(row["Liste rouge"])
        row_cells[2].text = str(row["Tronc"])

    document.add_page_break()

    annexe3 = document.add_paragraph('')
    annexe3.add_run(
        'Tableau 3 : ' + str(new_species_isole.shape[0]) + ' nouvelles espèces trouvées en ' + year + ' sur le tronc SMG du chemin du Pont-de-Bossy').bold = True

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Espèce'
    hdr_cells[1].text = 'Statut liste rouge'
    for index, row in new_species_isole.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Espèce"])
        if row["Espèce actuelle"] != row["Espèce"]:
            row_cells[0].text = str(row["Espèce"] + " (= " + row["Espèce actuelle"] + ")")
        row_cells[1].text = str(row["Liste rouge"])

    document.add_page_break()

    annexe4 = document.add_paragraph('')
    annexe4.add_run('Tableau 4 : ' + str(rare_species_maillettes.shape[0]) + ' espèces de la liste rouge, rares ou assez rares trouvées en ' + year + ' sur les troncs du chemin des Maillettes').bold = True

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Espèce'
    hdr_cells[1].text = 'Statut liste rouge'
    hdr_cells[2].text = 'Troncs'
    for index, row in rare_species_maillettes.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Espèce"])
        if row["Espèce actuelle"] != row["Espèce"]:
            row_cells[0].text = str(row["Espèce"] + " (= " + row["Espèce actuelle"] + ")")
        row_cells[1].text = str(row["Liste rouge"])
        row_cells[2].text = str(row["Tronc"])

    document.add_page_break()

    annexe5 = document.add_paragraph('')
    annexe5.add_run(
        'Tableau 5 : ' + str(rare_species_bossy.shape[0]) + ' espèces de la liste rouge, rares ou assez rares trouvées en ' + year + ' sur les troncs du chemin du Pont-de-Bossy').bold = True

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Espèce'
    hdr_cells[1].text = 'Statut liste rouge'
    hdr_cells[2].text = 'Troncs'
    for index, row in rare_species_bossy.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Espèce"])
        if row["Espèce actuelle"] != row["Espèce"]:
            row_cells[0].text = str(row["Espèce"] + " (= " + row["Espèce actuelle"] + ")")
        row_cells[1].text = str(row["Liste rouge"])
        row_cells[2].text = str(row["Tronc"])

    document.add_page_break()

    annexe6 = document.add_paragraph('')
    annexe6.add_run(
        'Tableau 6 : ' + str(rare_species_isole.shape[0]) + ' espèces de la liste rouge, rares ou assez rares trouvées en ' + year + ' sur le tronc SMG du chemin du Pont-de-Bossy').bold = True

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Espèce'
    hdr_cells[1].text = 'Statut liste rouge'
    for index, row in rare_species_isole.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Espèce"])
        if row["Espèce actuelle"] != row["Espèce"]:
            row_cells[0].text = str(row["Espèce"] + " (= " + row["Espèce actuelle"] + ")")
        row_cells[1].text = str(row["Liste rouge"])

    return document
